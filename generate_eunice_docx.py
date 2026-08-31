import os
import re
import io
import docx
from copy import deepcopy
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def append_doc_with_rebound_images_and_tables(src_path, dest_doc):
    src_doc = docx.Document(src_path)
    sectPr = dest_doc.element.body.xpath('w:sectPr')[0]
    
    for element in list(src_doc.element.body):
        tag = element.tag.split('}')[-1]
        
        if tag == 'p':
            drawings = element.xpath('.//w:drawing')
            if drawings:
                blips = element.xpath('.//a:blip')
                for b in blips:
                    rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and rId in src_doc.part.rels:
                        image_part = src_doc.part.rels[rId].target_part
                        image_bytes = image_part._blob
                        
                        p_img = dest_doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(14)
                        p_img.paragraph_format.space_after = Pt(4)
                        p_img.paragraph_format.keep_with_next = True
                        
                        run = p_img.add_run()
                        run.add_picture(io.BytesIO(image_bytes), width=Inches(6.2))
            else:
                p_src = None
                for p in src_doc.paragraphs:
                    if p._element == element:
                        p_src = p
                        break
                        
                if p_src and p_src.text.strip():
                    p_new = dest_doc.add_paragraph()
                    p_new.style = p_src.style
                    p_new.alignment = p_src.alignment
                    p_new.paragraph_format.space_before = p_src.paragraph_format.space_before
                    p_new.paragraph_format.space_after = p_src.paragraph_format.space_after
                    p_new.paragraph_format.line_spacing = p_src.paragraph_format.line_spacing
                    
                    for run in p_src.runs:
                        r_new = p_new.add_run(run.text)
                        r_new.bold = run.bold
                        r_new.italic = run.italic
                        r_new.font.name = 'Times New Roman'
                        if run.font.size:
                            r_new.font.size = run.font.size
                            
        elif tag == 'tbl':
            tbl_clone = deepcopy(element)
            sectPr.addprevious(tbl_clone)

def generate_combined_chapter_4_and_5():
    docs_dir = os.path.join(os.path.dirname(__file__), 'docs')
    ch4_docx = os.path.join(docs_dir, 'Chapter_4_Eunice.docx')
    ch5_docx = os.path.join(docs_dir, 'Chapter_5_Eunice.docx')
    combined_docx = os.path.join(docs_dir, 'Chapter_4_and_5_Eunice.docx')

    doc4 = docx.Document(ch4_docx)
    doc4.add_page_break()
    append_doc_with_rebound_images_and_tables(ch5_docx, doc4)
    doc4.save(combined_docx)
    print(f"Successfully generated Combined Chapter 4 & 5 Word Document at {combined_docx}")

def generate_full_dissertation():
    docs_dir = os.path.join(os.path.dirname(__file__), 'docs')
    ch123_docx = os.path.join(docs_dir, 'Chap_1_2_3_Eunice.docx')
    ch4_docx = os.path.join(docs_dir, 'Chapter_4_Eunice.docx')
    ch5_docx = os.path.join(docs_dir, 'Chapter_5_Eunice.docx')
    full_docx = os.path.join(docs_dir, 'Chapter_1_2_3_4_5_Eunice.docx')

    doc123 = docx.Document(ch123_docx)
    doc123.add_page_break()
    append_doc_with_rebound_images_and_tables(ch4_docx, doc123)
    doc123.add_page_break()
    append_doc_with_rebound_images_and_tables(ch5_docx, doc123)

    doc123.save(full_docx)
    print(f"Successfully generated Full Dissertation Word Document at {full_docx}")

if __name__ == "__main__":
    generate_combined_chapter_4_and_5()
    generate_full_dissertation()
